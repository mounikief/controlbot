"""
Report Builder für ControlBot
Erstellt professionelle Word-Dokumente aus Report-Daten
"""

import os
from datetime import datetime
from typing import Dict, Any
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import io
from PIL import Image

class ReportBuilder:
    """Klasse für Erstellung von Word-Reports"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Richtet Custom Styles für das Dokument ein"""
        styles = self.doc.styles
        
        # Titel-Style
        if 'CustomTitle' not in styles:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(31, 119, 180)  # Blau
        
        # Heading-Style
        if 'CustomHeading' not in styles:
            heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Calibri'
            heading_font.size = Pt(16)
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(68, 114, 196)
        
        # Subheading-Style
        if 'CustomSubheading' not in styles:
            subheading_style = styles.add_style('CustomSubheading', WD_STYLE_TYPE.PARAGRAPH)
            subheading_font = subheading_style.font
            subheading_font.name = 'Calibri'
            subheading_font.size = Pt(12)
            subheading_font.bold = True
    
    def create_word_report(
        self,
        report_content: str,
        df: pd.DataFrame,
        include_charts: bool = True
    ) -> str:
        """
        Erstellt ein Word-Dokument aus Report-Content
        
        Args:
            report_content: Markdown-formatierter Report-Text
            df: DataFrame mit Projektdaten
            include_charts: Ob Diagramme eingebettet werden sollen
            
        Returns:
            Pfad zum erstellten Word-Dokument
        """
        # Neues Dokument
        self.doc = Document()
        self._setup_styles()
        
        # Füge Header hinzu
        self._add_header()
        
        # Füge Report-Content hinzu
        self._add_content_from_markdown(report_content)
        
        # Füge Datentabelle hinzu
        self._add_data_table(df)
        
        # Füge Charts hinzu (wenn gewünscht)
        if include_charts:
            try:
                self._add_charts_section(df)
            except Exception as e:
                print(f"Warnung: Charts konnten nicht hinzugefügt werden: {str(e)}")
        
        # Füge Footer hinzu
        self._add_footer()
        
        # Speichere Dokument
        output_path = f"/home/claude/controlbot_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        self.doc.save(output_path)
        
        return output_path
    
    def _add_header(self):
        """Fügt den Dokumenten-Header hinzu"""
        # Logo-Platzhalter (optional)
        title = self.doc.add_paragraph('ControlBot')
        title.style = 'CustomTitle'
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = self.doc.add_paragraph('Projektcontrolling Report')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.runs[0]
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Datum
        date_para = self.doc.add_paragraph(f"Erstellt am: {datetime.now().strftime('%d.%m.%Y um %H:%M Uhr')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.runs[0]
        run.font.size = Pt(10)
        run.font.italic = True
        
        self.doc.add_paragraph()  # Leerzeile
        self.doc.add_paragraph('─' * 80)  # Trennlinie
        self.doc.add_paragraph()
    
    def _add_content_from_markdown(self, markdown_text: str):
        """
        Konvertiert Markdown-Text in Word-Formatierung
        
        Args:
            markdown_text: Markdown-formatierter Text
        """
        lines = markdown_text.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line or line == '---':
                continue
            
            # Hauptüberschriften (# )
            if line.startswith('# '):
                para = self.doc.add_paragraph(line[2:])
                para.style = 'CustomTitle'
            
            # Überschriften (## )
            elif line.startswith('## '):
                self.doc.add_paragraph()  # Abstand vor Überschrift
                para = self.doc.add_paragraph(line[3:])
                para.style = 'CustomHeading'
            
            # Unterüberschriften (### )
            elif line.startswith('### '):
                para = self.doc.add_paragraph(line[4:])
                para.style = 'CustomSubheading'
            
            # Bullet Points (- oder *)
            elif line.startswith('- ') or line.startswith('* '):
                para = self.doc.add_paragraph(line[2:], style='List Bullet')
            
            # Nummerierte Listen (1. )
            elif len(line) > 2 and line[0].isdigit() and line[1:3] == '. ':
                para = self.doc.add_paragraph(line[3:], style='List Number')
            
            # Fett gedruckt (**text**)
            elif '**' in line:
                para = self.doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = para.add_run(part)
                    if i % 2 == 1:  # Ungerade Index = zwischen **
                        run.bold = True
            
            # Normaler Text
            else:
                if line:  # Nur nicht-leere Zeilen
                    self.doc.add_paragraph(line)
    
    def _add_data_table(self, df: pd.DataFrame):
        """
        Fügt eine formatierte Tabelle mit Projektdaten hinzu
        
        Args:
            df: DataFrame mit Projektdaten
        """
        self.doc.add_page_break()
        
        heading = self.doc.add_paragraph('Anhang: Detaillierte Projektdaten')
        heading.style = 'CustomHeading'
        
        # Bereite Daten vor
        display_columns = ['projekt_name', 'kosten_plan', 'kosten_ist']
        if 'kosten_abweichung_prozent' in df.columns:
            display_columns.append('kosten_abweichung_prozent')
        if 'kosten_status' in df.columns:
            display_columns.append('kosten_status')
        
        df_display = df[display_columns].copy()
        
        # Erstelle Tabelle
        table = self.doc.add_table(rows=1, cols=len(display_columns))
        table.style = 'Light Grid Accent 1'
        
        # Header-Zeile
        header_cells = table.rows[0].cells
        headers = {
            'projekt_name': 'Projektname',
            'kosten_plan': 'Plan (€)',
            'kosten_ist': 'Ist (€)',
            'kosten_abweichung_prozent': 'Abweichung (%)',
            'kosten_status': 'Status'
        }
        
        for i, col in enumerate(display_columns):
            cell = header_cells[i]
            cell.text = headers.get(col, col)
            # Fett für Header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        
        # Daten-Zeilen
        for _, row in df_display.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(display_columns):
                value = row[col]
                
                # Formatierung je nach Spalte
                if col == 'kosten_plan' or col == 'kosten_ist':
                    row_cells[i].text = f"€{value:,.0f}"
                elif col == 'kosten_abweichung_prozent':
                    row_cells[i].text = f"{value:.1f}%"
                    # Farbcodierung
                    if value > 10:
                        for paragraph in row_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(192, 0, 0)  # Rot
                    elif value > 5:
                        for paragraph in row_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(255, 153, 0)  # Orange
                else:
                    row_cells[i].text = str(value)
                
                # Kleinere Schrift für Daten
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    
    def _add_charts_section(self, df: pd.DataFrame):
        """
        Fügt eine Sektion mit Charts hinzu
        
        Args:
            df: DataFrame mit Projektdaten
        """
        try:
            import matplotlib.pyplot as plt
            import matplotlib
            matplotlib.use('Agg')  # Non-interactive backend
            
            self.doc.add_page_break()
            
            heading = self.doc.add_paragraph('Visualisierungen')
            heading.style = 'CustomHeading'
            
            # Chart 1: Kosten Vergleich
            fig, ax = plt.subplots(figsize=(10, 6))
            
            projects = df['projekt_name'][:10]  # Top 10 Projekte
            x_pos = range(len(projects))
            
            plan_values = df['kosten_plan'][:10]
            ist_values = df['kosten_ist'][:10]
            
            width = 0.35
            ax.bar([x - width/2 for x in x_pos], plan_values, width, label='Plan', color='lightblue')
            ax.bar([x + width/2 for x in x_pos], ist_values, width, label='Ist', color='coral')
            
            ax.set_xlabel('Projekt')
            ax.set_ylabel('Kosten (€)')
            ax.set_title('Kosten: Plan vs. Ist (Top 10 Projekte)')
            ax.set_xticks(x_pos)
            ax.set_xticklabels(projects, rotation=45, ha='right')
            ax.legend()
            ax.grid(True, alpha=0.3)
            
            plt.tight_layout()
            
            # Speichere Chart als Bild
            chart_path = '/home/claude/temp_chart1.png'
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close()
            
            # Füge Bild zum Dokument hinzu
            self.doc.add_picture(chart_path, width=Inches(6))
            
            # Lösche temporäre Datei
            if os.path.exists(chart_path):
                os.remove(chart_path)
            
            # Chart 2: Abweichungsanalyse
            if 'kosten_abweichung_prozent' in df.columns:
                self.doc.add_paragraph()
                
                fig, ax = plt.subplots(figsize=(10, 6))
                
                abweichungen = df['kosten_abweichung_prozent'][:10]
                colors = ['red' if x > 10 else 'orange' if x > 0 else 'green' for x in abweichungen]
                
                ax.barh(projects, abweichungen, color=colors)
                ax.set_xlabel('Abweichung (%)')
                ax.set_title('Kostenabweichungen (Top 10 Projekte)')
                ax.axvline(x=0, color='black', linestyle='-', linewidth=0.8)
                ax.axvline(x=10, color='red', linestyle='--', linewidth=0.5, alpha=0.7)
                ax.grid(True, alpha=0.3)
                
                plt.tight_layout()
                
                chart_path = '/home/claude/temp_chart2.png'
                plt.savefig(chart_path, dpi=150, bbox_inches='tight')
                plt.close()
                
                self.doc.add_picture(chart_path, width=Inches(6))
                
                if os.path.exists(chart_path):
                    os.remove(chart_path)
        
        except ImportError:
            self.doc.add_paragraph("Hinweis: Charts konnten nicht erstellt werden (matplotlib nicht installiert)")
        except Exception as e:
            self.doc.add_paragraph(f"Hinweis: Charts konnten nicht erstellt werden ({str(e)})")
    
    def _add_footer(self):
        """Fügt den Dokumenten-Footer hinzu"""
        self.doc.add_paragraph()
        self.doc.add_paragraph('─' * 80)
        
        footer = self.doc.add_paragraph()
        run = footer.add_run('Dieser Report wurde automatisch mit ControlBot generiert.\n')
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        run = footer.add_run(f'© {datetime.now().year} ControlBot | www.controlbot.de')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

if __name__ == "__main__":
    # Test des Moduls
    print("Report Builder Module geladen.")
    
    # Erstelle Test-Report
    builder = ReportBuilder()
    
    test_content = """
# Test Report

## Executive Summary
Dies ist ein Test-Report.

## Wichtige Erkenntnisse
- Punkt 1
- Punkt 2
- Punkt 3

**Fazit:** Alles funktioniert!
"""
    
    test_df = pd.DataFrame({
        'projekt_name': ['Projekt A', 'Projekt B'],
        'kosten_plan': [100000, 200000],
        'kosten_ist': [110000, 190000],
        'kosten_abweichung_prozent': [10, -5],
        'kosten_status': ['Warnung', 'Im Plan']
    })
    
    output_path = builder.create_word_report(test_content, test_df, include_charts=False)
    print(f"Test-Report erstellt: {output_path}")
